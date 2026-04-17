"""
converter.py — Centro de Conversión Universal
Macro-Módulo 2: Convierte entre formatos de documentos (PDF, TXT, DOCX, Imágenes)
y formatos de datos (CSV, Excel, JSON, Parquet). Toda conversión de datos usa chunks.
"""

import os
import json
from typing import Optional, Callable, List

import pandas as pd

from src.config.settings import CHUNK_EXCEL
from src.utils.pdf_utils import compress_pdf, unir_pdfs


# ─────────────────────────────────────────────
# 2.1 Conversor de Documentos (estilo iLovePDF)
# ─────────────────────────────────────────────

def txt_a_pdf(ruta_txt: str, ruta_pdf: str,
              callback: Optional[Callable] = None) -> str:
    """Convierte un archivo TXT a PDF."""
    from fpdf import FPDF
    if callback:
        callback("Convirtiendo TXT → PDF...")

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Courier", size=10)

    with open(ruta_txt, 'r', encoding='utf-8', errors='replace') as f:
        for line in f:
            line_clean = line.encode('latin-1', 'replace').decode('latin-1')
            pdf.multi_cell(0, 5, line_clean)

    pdf.output(ruta_pdf)
    ruta_pdf = compress_pdf(ruta_pdf, callback=callback)
    return ruta_pdf


def pdf_a_txt(ruta_pdf: str, ruta_txt: str,
              callback: Optional[Callable] = None) -> str:
    """Extrae el texto de un PDF y lo guarda como TXT."""
    from pypdf import PdfReader
    if callback:
        callback("Convirtiendo PDF → TXT...")

    reader = PdfReader(ruta_pdf)
    with open(ruta_txt, 'w', encoding='utf-8') as f:
        for i, page in enumerate(reader.pages):
            texto = page.extract_text()
            if texto:
                f.write(f"\n--- PÁGINA {i+1} ---\n")
                f.write(texto)
            if callback and (i + 1) % 5 == 0:
                callback(f"   Página {i+1} procesada...")

    return ruta_txt


def imagenes_a_pdf(rutas_imagenes: List[str], ruta_pdf: str,
                   callback: Optional[Callable] = None) -> str:
    """Convierte una o varias imágenes a un único PDF."""
    from PIL import Image
    if callback:
        callback(f"Convirtiendo {len(rutas_imagenes)} imagen(es) → PDF...")

    if not rutas_imagenes:
        raise ValueError("No se proporcionaron imágenes.")

    imgs = []
    for p in rutas_imagenes:
        img = Image.open(p)
        if img.mode != 'RGB':
            img = img.convert('RGB')
        imgs.append(img)

    primera, resto = imgs[0], imgs[1:]
    primera.save(ruta_pdf, save_all=True, append_images=resto)

    for img in imgs:
        img.close()

    return ruta_pdf


def txt_a_docx(ruta_txt: str, ruta_docx: str,
               callback: Optional[Callable] = None) -> str:
    """Convierte un archivo TXT a DOCX (Microsoft Word)."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("Instala python-docx: pip install python-docx")

    if callback:
        callback("Convirtiendo TXT → DOCX...")

    doc = Document()
    doc.add_heading('Documento Convertido', level=1)

    with open(ruta_txt, 'r', encoding='utf-8', errors='replace') as f:
        for line in f:
            doc.add_paragraph(line.rstrip())

    doc.save(ruta_docx)
    return ruta_docx


def pdf_a_docx(ruta_pdf: str, ruta_docx: str,
               callback: Optional[Callable] = None) -> str:
    """Extrae texto de PDF y lo guarda como DOCX."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("Instala python-docx: pip install python-docx")
    from pypdf import PdfReader

    if callback:
        callback("Convirtiendo PDF → DOCX...")

    reader = PdfReader(ruta_pdf)
    doc = Document()
    doc.add_heading('Documento Exportado desde PDF', level=1)

    for i, page in enumerate(reader.pages):
        texto = page.extract_text()
        if texto:
            doc.add_heading(f'Página {i+1}', level=2)
            doc.add_paragraph(texto)

    doc.save(ruta_docx)
    return ruta_docx


# ─────────────────────────────────────────────
# 2.2 Conversor de Datos — ETL Lite (chunks)
# ─────────────────────────────────────────────

def csv_a_excel(ruta_csv: str, ruta_xlsx: str,
                callback: Optional[Callable] = None) -> str:
    """Convierte CSV a Excel usando chunks para no saturar RAM."""
    if callback:
        callback("Convirtiendo CSV → Excel (chunks)...")

    reader = pd.read_csv(ruta_csv, chunksize=CHUNK_EXCEL, low_memory=False)
    with pd.ExcelWriter(ruta_xlsx, engine='openpyxl') as writer:
        for i, chunk in enumerate(reader):
            chunk.to_excel(
                writer, sheet_name='Data', index=False,
                startrow=i * CHUNK_EXCEL, header=(i == 0)
            )
            if callback:
                callback(f"   Bloque {i+1} escrito...")

    return ruta_xlsx


def excel_a_csv(ruta_xlsx: str, ruta_csv: str,
                callback: Optional[Callable] = None) -> str:
    """Convierte Excel a CSV."""
    if callback:
        callback("Convirtiendo Excel → CSV...")
    df = pd.read_excel(ruta_xlsx)
    df.to_csv(ruta_csv, index=False)
    return ruta_csv


def csv_a_json(ruta_csv: str, ruta_json: str,
               callback: Optional[Callable] = None) -> str:
    """Convierte CSV a JSON (línea a línea para archivos grandes)."""
    if callback:
        callback("Convirtiendo CSV → JSON (chunks)...")

    registros = []
    for chunk in pd.read_csv(ruta_csv, chunksize=CHUNK_EXCEL, low_memory=False):
        registros.extend(chunk.to_dict(orient='records'))
        if callback:
            callback(f"   {len(registros):,} registros procesados...")

    with open(ruta_json, 'w', encoding='utf-8') as f:
        json.dump(registros, f, ensure_ascii=False, indent=2, default=str)

    return ruta_json


def json_a_csv(ruta_json: str, ruta_csv: str,
               callback: Optional[Callable] = None) -> str:
    """Convierte JSON a CSV."""
    if callback:
        callback("Convirtiendo JSON → CSV...")
    with open(ruta_json, 'r', encoding='utf-8') as f:
        data = json.load(f)
    df = pd.DataFrame(data if isinstance(data, list) else [data])
    df.to_csv(ruta_csv, index=False)
    return ruta_csv


def excel_a_parquet(ruta_xlsx: str, ruta_parquet: str,
                    callback: Optional[Callable] = None) -> str:
    """Convierte Excel a Parquet (formato columnar eficiente)."""
    if callback:
        callback("Convirtiendo Excel → Parquet...")
    df = pd.read_excel(ruta_xlsx)
    df.to_parquet(ruta_parquet, index=False)
    if callback:
        callback(f"✅ Parquet guardado: {os.path.getsize(ruta_parquet)/1024:.1f} KB")
    return ruta_parquet


def parquet_a_excel(ruta_parquet: str, ruta_xlsx: str,
                    callback: Optional[Callable] = None) -> str:
    """Convierte Parquet a Excel usando chunks."""
    if callback:
        callback("Convirtiendo Parquet → Excel (chunks)...")

    import pyarrow.parquet as pq
    pf = pq.ParquetFile(ruta_parquet)

    with pd.ExcelWriter(ruta_xlsx, engine='openpyxl') as writer:
        fila_inicio = 0
        primer_chunk = True
        for batch in pf.iter_batches(batch_size=CHUNK_EXCEL):
            chunk = batch.to_pandas()
            chunk.to_excel(writer, sheet_name='Data', index=False,
                           startrow=fila_inicio, header=primer_chunk)
            fila_inicio += len(chunk) + (1 if primer_chunk else 0)
            primer_chunk = False
            if callback:
                callback(f"   {fila_inicio:,} filas escritas...")

    return ruta_xlsx


def csv_a_parquet(ruta_csv: str, ruta_parquet: str,
                  callback: Optional[Callable] = None) -> str:
    """Convierte CSV a Parquet."""
    if callback:
        callback("Convirtiendo CSV → Parquet...")
    df = pd.read_csv(ruta_csv, low_memory=False)
    df.to_parquet(ruta_parquet, index=False)
    return ruta_parquet


def parquet_a_csv(ruta_parquet: str, ruta_csv: str,
                  callback: Optional[Callable] = None) -> str:
    """Convierte Parquet a CSV."""
    if callback:
        callback("Convirtiendo Parquet → CSV...")
    import pyarrow.parquet as pq
    pf = pq.ParquetFile(ruta_parquet)
    primer = True
    with open(ruta_csv, 'w', encoding='utf-8', newline='') as f:
        for batch in pf.iter_batches(batch_size=CHUNK_EXCEL):
            batch.to_pandas().to_csv(f, index=False, header=primer)
            primer = False
    return ruta_csv


# ─────────────────────────────────────────────
# Función maestra de enrutamiento
# ─────────────────────────────────────────────

# Mapa de conversiones soportadas: (ext_origen, ext_destino) → función
_RUTAS: dict = {
    ('.csv',   'xlsx'):    csv_a_excel,
    ('.csv',   'json'):    csv_a_json,
    ('.csv',   'parquet'): csv_a_parquet,
    ('.xlsx',  'csv'):     excel_a_csv,
    ('.xls',   'csv'):     excel_a_csv,
    ('.xlsx',  'parquet'): excel_a_parquet,
    ('.xls',   'parquet'): excel_a_parquet,
    ('.parquet','csv'):    parquet_a_csv,
    ('.parquet','xlsx'):   parquet_a_excel,
    ('.pq',    'csv'):     parquet_a_csv,
    ('.pq',    'xlsx'):    parquet_a_excel,
    ('.json',  'csv'):     json_a_csv,
    ('.txt',   'pdf'):     txt_a_pdf,
    ('.txt',   'docx'):    txt_a_docx,
    ('.pdf',   'txt'):     pdf_a_txt,
    ('.pdf',   'docx'):    pdf_a_docx,
    ('.png',   'pdf'):     lambda o, d, cb=None: imagenes_a_pdf([o], d, cb),
    ('.jpg',   'pdf'):     lambda o, d, cb=None: imagenes_a_pdf([o], d, cb),
    ('.jpeg',  'pdf'):     lambda o, d, cb=None: imagenes_a_pdf([o], d, cb),
}


def convertir_archivo(
    ruta_origen: str,
    formato_destino: str,
    carpeta_salida: str,
    callback: Optional[Callable] = None,
) -> str:
    """
    Convierte un archivo al formato de destino especificado.

    Args:
        ruta_origen:      Ruta del archivo de origen.
        formato_destino:  Extensión del formato destino (ej: 'xlsx', 'txt').
        carpeta_salida:   Carpeta donde se guardará el archivo convertido.
        callback:         Función de progreso opcional.

    Returns:
        Ruta del archivo convertido.

    Raises:
        ValueError: Si la conversión no está soportada.
    """
    nombre_base = os.path.splitext(os.path.basename(ruta_origen))[0]
    ext_origen = os.path.splitext(ruta_origen)[1].lower()
    fmt_dest = formato_destino.lower().lstrip('.')

    ruta_destino = os.path.join(carpeta_salida, f"{nombre_base}.{fmt_dest}")

    clave = (ext_origen, fmt_dest)
    funcion = _RUTAS.get(clave)

    if funcion is None:
        raise ValueError(
            f"Conversión de '{ext_origen}' → '{fmt_dest}' no soportada.\n"
            f"Conversiones disponibles: {sorted(set(k[0] for k in _RUTAS))} → "
            f"{sorted(set(k[1] for k in _RUTAS))}"
        )

    try:
        os.makedirs(carpeta_salida, exist_ok=True)
        return funcion(ruta_origen, ruta_destino, callback)
    except Exception as exc:
        if callback:
            callback(f"❌ Error en conversión: {exc}")
        raise


def unir_pdfs_wrapper(
    lista_rutas: List[str],
    ruta_salida: str,
    callback: Optional[Callable] = None,
) -> str:
    """Combina múltiples PDFs en uno solo."""
    return unir_pdfs(lista_rutas, ruta_salida, callback)


def formatos_soportados_para(ext_origen: str) -> List[str]:
    """Retorna los formatos de destino disponibles para una extensión dada."""
    ext = ext_origen.lower()
    return [dest for (orig, dest) in _RUTAS.keys() if orig == ext]
